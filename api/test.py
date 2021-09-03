import shutil
import os

from flask import Flask , request,jsonify, json,current_app
from docx import Document
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn,nsdecls
from datetime import datetime
from flask_sqlalchemy import SQLAlchemy
import schedule


# app = Flask(__name__)


# app.config['SQLALCHEMY_DATABASE_URI']="mysql://earlyfont:earlyfont@localhost:3306/earlyfont"
# app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] =False

# db=SQLAlchemy(app)

def refresh_estimate():
    file_names = ['estimate_basic.docx','estimate_basicPlus.docx','estimate_premium.docx','estimate_premiumPlus.docx']
    # for idx,name in enumerate(file_names):
    name = 'estimate_basic.docx'
    print(os. getcwd())
    path = os.path.join(os. getcwd(),'../public/Estimates')
    shutil.copy(
    name,  os.path.join(path, name))
    document = Document("{}/{}".format(path, name))
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"견적일 {datetime.today().strftime('%Y.%m.%d')} / 견적일로부터 15일간 유효합니다.")
    run.font.name = "에스코어 드림 5 Medium"
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '에스코어 드림 5 Medium')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(os.path.join(path,name))
    convert(os.path.join(path,name))
refresh_estimate()
