import shutil
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn,nsdecls
from flask import Flask , request,jsonify, json,current_app
from datetime import datetime
from flask_sqlalchemy import SQLAlchemy
from flask_apscheduler import APScheduler
import pythoncom
import win32com.client as client


app = Flask(__name__,static_folder='../build', static_url_path='/')
app.app_context()

app.config['SQLALCHEMY_DATABASE_URI']="mysql://earlyfont:earlyfont@localhost:3306/earlyfont"
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] =False

db=SQLAlchemy(app)
scheduler = APScheduler()

class Uploads(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    date = db.Column(db.DateTime, nullable=True, default=datetime.now)
    title = db.Column(db.String(100), nullable=False)
    thumnail = db.Column(db.String(100), nullable=False)
    detail_image = db.Column(db.String(100), nullable=False)
    font_file = db.Column(db.Boolean, nullable=True, default=False)
    
    def serializer(font_info):
        return {
            'id': font_info.id,
            'date': font_info.date.strftime("%Y/%m/%d %H:%M"),
            'title': font_info.title,
            'thumnail' : font_info.thumnail,
            'detail_image': font_info.detail_image,
            'font_file' : font_info.font_file
        }
        
    def __repr__(self):
        return f"Uploads(title : {self.title},date : {self.date})"

@app.route('/admin')
@app.route('/branding')
@app.route('/portfolio')
@app.route('/estimate')
@app.route('/')
def index():
    return current_app.send_static_file('index.html')

@app.route('/portfolio/<string:name>')
def index2(name):
    return current_app.send_static_file('index.html')

@app.route('/sitemap.xml')
def sitemap():
    return current_app.send_static_file('sitemap.xml')


def save_file(file,file_name):
    path = os.path.join(
        current_app.root_path,'../build/FontImages',file_name)
    path2 = os.path.join(
        current_app.root_path,'../public/FontImages',file_name)
    file.save(path)
    file.save(path2)
    
    
@app.route('/api/upload',methods=['GET','POST'])
def upload():
    num = Uploads.query.order_by(Uploads.id.desc()).first().id+1
    title = request.form.get('title')
    thumnail = f"/FontImages/{num}_1.jpg"
    detail_image = f"/FontImages/{num}_2.jpg"
    save_file(request.files.get(f"file1"),f"{num}_1.jpg")
    save_file(request.files.get(f"file2"),f"{num}_2.jpg")
    new_font= Uploads(title=title,thumnail=thumnail,detail_image=detail_image)
    if request.files.get(f"file3"):
        save_file(request.files.get(f"file3"),f"{title}.zip")
        new_font= Uploads(title=title,thumnail=thumnail,detail_image=detail_image,font_file=True)
    db.session.add(new_font)
    db.session.commit()

    return jsonify({'success':'등록되었습니다!'})

@app.route('/api/edit',methods=['GET','POST'])
def edit():
    fontInfo = Uploads.query.filter_by(id=request.form.get('id')).first()
    title = request.form.get('title')
    thumnail = f"/FontImages/{fontInfo.id}_1.jpg"
    if request.files.get(f"file1"):
        save_file(request.files.get(f"file1"),f"{fontInfo.id}_1.jpg")
        fontInfo.thumnail = thumnail
    detail_image = f"/FontImages/{fontInfo.id}_2.jpg"
    if request.files.get(f"file2"):
        save_file(request.files.get(f"file2"),f"{fontInfo.id}_2.jpg")
        fontInfo.detail_image = detail_image
    if request.files.get(f"file3") :
        save_file(request.files.get(f"file3"),f"{title}.zip")
        fontInfo.font_file = True
    fontInfo.title = title
    db.session.commit()

    return jsonify({'success':'변경되었습니다!'})

@app.route('/api/fontList',methods=['GET','POST'])
def fontList():
    lists = Uploads.query.order_by(Uploads.id.desc()).all()
    return jsonify({'uploaded_list':[*map(Uploads.serializer,lists)]})

@app.route('/api/delete',methods=['GET','POST'])
def delete():
    request_data = json.loads(request.data)
    selected_font = Uploads.query.filter_by(id=request_data['id']).first()
    db.session.delete(selected_font)
    db.session.commit()
    
    return jsonify({'success':'삭제하였습니다.'})
    
@app.route('/api/loadFontData',methods=['GET','POST'])
def loadFontData():
    request_data = json.loads(request.data)
    font = Uploads.query.filter_by(id=int(request_data['id'])).first()
    return jsonify({'font': Uploads.serializer(font)})

@app.route('/api/loadFontSingleData',methods=['GET','POST'])
def loadFontSingleData():
    request_data = json.loads(request.data)
    font = Uploads.query.filter_by(title=request_data['name']).first()
    return jsonify({'font': Uploads.serializer(font)})

def convert_to_pdf(doc):
    try:
        word = client.DispatchEx("Word.Application")
        new_name = doc.replace(".docx", r".pdf")
        worddoc = word.Documents.Open(doc)
        worddoc.SaveAs(new_name, FileFormat = 17)
        worddoc.Close()
    except Exception as e:
            return e
    finally:
            word.Quit()
            
def refresh_estimate():
    pythoncom.CoInitialize()
    file_names = ['Earlyfont_basic.docx','Earlyfont_basicPlus.docx','Earlyfont_premium.docx','Earlyfont_premiumPlus.docx']
    for idx,name in enumerate(file_names):
        path = os.path.join(current_app.root_path,'../build/Estimates')
        shutil.copy(
        name,  os.path.join(path, name))
        document = Document("{}/{}".format(path, name))
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"견적일 {datetime.today().strftime('%Y.%m.%d')} / 견적일로부터 15일간 유효합니다.")
        run.font.name = "에스코어 드림 5 Medium"
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '에스코어 드림 5 Medium')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.save(os.path.join(path,name))
        convert_to_pdf(os.path.join(path,name))


            
def daily_refresh():
    with app.app_context():
        refresh_estimate()






    
if __name__ =='__main__':
    scheduler.add_job(id='daily update estimates',func=daily_refresh,trigger='cron',hour=0)
    scheduler.start()
    app.run(debug=True,port=8080)
    

    
